"""Document loaders: turn a file on disk into plain text.

Each format has its own small, directly testable function (:func:`load_text`,
:func:`load_pdf`). :func:`load` is a dispatcher that picks the right one based
on the file extension, so callers don't need to care about the format.
"""

import logging
from collections.abc import Callable
from pathlib import Path
from typing import Any

import docx
import pypdf

from industryiq.config import get_settings

logger = logging.getLogger(__name__)


def load_text(path: str | Path) -> str:
    """Read a plain-text ``.txt`` file and return its contents (UTF-8).

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.txt`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".txt":
        raise ValueError(f"load_text expects a .txt file, got {p.suffix!r}")
    return p.read_text(encoding="utf-8")


def load_pdf_pages(path: str | Path) -> list[str]:
    """Extract a ``.pdf`` file's text page by page (one string per page).

    Which engine does the extraction is set by ``PDF_PARSER``:

    * ``"docling"`` (default) -- layout-aware parsing that emits Markdown with
      correct reading order and headings, which chunks/retrieves far better on
      multi-column report PDFs. Slower (seconds/page) and needs the optional
      ``docling`` extra (``pip install 'industryiq[docling]'``). If Docling fails
      for any reason -- not installed, or a PDF it can't convert -- this falls
      back to pypdf so a long offline ingest is never halted by one bad file.
    * ``"pypdf"`` -- fast, pure-Python text extraction with no fallback. Fine for
      clean single-column PDFs; weak on multi-column layouts.

    Either way the result is one string per page, so page-number citations work.

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.pdf`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".pdf":
        raise ValueError(f"load_pdf expects a .pdf file, got {p.suffix!r}")
    if get_settings().pdf_parser == "docling":
        try:
            return _load_pdf_pages_docling(p)
        except Exception as exc:  # noqa: BLE001 -- any Docling failure falls back to pypdf
            logger.warning("Docling failed on %s (%s); falling back to pypdf.", p.name, exc)
            return _load_pdf_pages_pypdf(p)
    return _load_pdf_pages_pypdf(p)


def _load_pdf_pages_pypdf(p: Path) -> list[str]:
    """Per-page plain text via pypdf (the default engine)."""
    reader = pypdf.PdfReader(str(p))
    return [page.extract_text() for page in reader.pages]


# Building a Docling converter loads ML models, so build it once and reuse it.
_docling_converter: Any = None


def _patch_rapidocr_scale(scale: int) -> None:
    """Lower the resolution at which Docling renders page regions for OCR.

    Docling's RapidOCR stage hardcodes ``self.scale = 3`` (216 DPI, then x1.5 =
    324 DPI) and exposes no option to change it, so those high-res renders pile up
    and OOM/SIGSEGV the process on large reports. We patch the model class once to
    set a lower scale. Best-effort: if Docling's internals have shifted, we log and
    leave the default rather than break ingestion.
    """
    try:  # pragma: no cover - needs the heavy extra
        from docling.models.stages.ocr import rapid_ocr_model

        if getattr(rapid_ocr_model.RapidOcrModel, "_iiq_scale_patched", False):
            return
        _orig_init = rapid_ocr_model.RapidOcrModel.__init__

        def _scaled_init(self: Any, *args: Any, **kwargs: Any) -> None:
            _orig_init(self, *args, **kwargs)
            self.scale = scale

        rapid_ocr_model.RapidOcrModel.__init__ = _scaled_init
        rapid_ocr_model.RapidOcrModel._iiq_scale_patched = True
    except Exception as exc:  # noqa: BLE001 -- Docling internals can shift across versions
        logger.warning("Could not lower RapidOCR render scale (%s); using its default.", exc)


def _get_docling_converter() -> Any:
    """Return a cached Docling ``DocumentConverter``, built on first use.

    OCR is on by default (``DOCLING_OCR``). RapidOCR's detection step is forced to
    limit_type=max so a large embedded bitmap is downscaled (to RapidOCR's internal
    2000px ceiling) before inference -- its default (limit_type=min) only upscales,
    so a full-size chart bitmap grows the ONNX tensor until it OOMs
    (``std::bad_alloc``). Raises a pointed error if the optional ``docling`` extra
    isn't installed, rather than the bare ``ModuleNotFoundError`` the import gives.
    """
    global _docling_converter
    if _docling_converter is None:
        try:
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions, RapidOcrOptions
            from docling.datamodel.settings import settings as docling_settings
            from docling.document_converter import DocumentConverter, PdfFormatOption
        except ModuleNotFoundError as exc:
            raise ModuleNotFoundError(
                "PDF_PARSER='docling' needs the optional 'docling' dependency; "
                "install it with:  pip install 'industryiq[docling]'"
            ) from exc
        settings = get_settings()  # pragma: no cover - needs the heavy extra
        # Serialize page rasterization to cap peak memory; the default (4 pages at
        # once) can OOM the whole page on large media, dropping its text too.
        docling_settings.perf.page_batch_size = settings.docling_page_batch_size
        if settings.docling_ocr:
            _patch_rapidocr_scale(settings.docling_ocr_scale)
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = settings.docling_ocr
        # Force the detection step to downscale large bitmaps; RapidOCR's default
        # (limit_type=min) never shrinks them, so a full-size chart bitmap OOMs.
        pipeline_options.ocr_options = RapidOcrOptions(rapidocr_params={"Det.limit_type": "max"})
        _docling_converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
        )
    return _docling_converter


def _load_pdf_pages_docling(p: Path) -> list[str]:
    """Per-page Markdown via Docling (layout-aware, opt-in engine).

    Falls back to a single whole-document element if Docling reports no pages.
    """
    doc: Any = _get_docling_converter().convert(str(p)).document
    page_count = len(doc.pages)
    if page_count == 0:
        return [doc.export_to_markdown()]
    return [doc.export_to_markdown(page_no=n) for n in range(1, page_count + 1)]


def load_pdf(path: str | Path) -> str:
    """Extract text from a ``.pdf`` file, joining pages with newlines."""
    return "\n".join(load_pdf_pages(path))


def load_docx(path: str | Path) -> str:
    """Extract text from a ``.docx`` file, joining paragraphs with newlines.

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file is not a ``.docx`` file.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"load_docx expects a .docx file, got {p.suffix!r}")
    document = docx.Document(str(p))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


_LOADERS: dict[str, Callable[[str | Path], str]] = {
    ".txt": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
}

SUPPORTED_EXTENSIONS = frozenset(_LOADERS)


def _to_utf8_safe(text: str) -> str:
    """Drop characters that cannot be encoded as UTF-8.

    PDF extraction can emit lone surrogate code points (from broken font maps)
    that are valid in a Python ``str`` but not encodable to UTF-8. Left in, they
    crash every UTF-8 consumer downstream -- the embedding tokenizer, JSON
    payloads to Bedrock, and Postgres text columns. Stripping them here keeps
    each loader's output safe to embed and store.
    """
    return text.encode("utf-8", "ignore").decode("utf-8")


def load(path: str | Path) -> str:
    """Load any supported file by dispatching on its extension.

    The returned text is guaranteed UTF-8 encodable (see :func:`_to_utf8_safe`).

    Raises:
        FileNotFoundError: If ``path`` does not point to an existing file.
        ValueError: If the file's extension is not supported.
    """
    p = Path(path)
    loader = _LOADERS.get(p.suffix.lower())
    if loader is None:
        raise ValueError(
            f"Unsupported file type {p.suffix!r}; supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return _to_utf8_safe(loader(path))


def load_pages(path: str | Path) -> list[str]:
    """Load a file as a list of page texts (each UTF-8 safe), for page attribution.

    PDFs return one element per page; other formats have no real pagination, so
    they return the whole document as a single element.
    """
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return [_to_utf8_safe(page) for page in load_pdf_pages(p)]
    return [load(p)]


def load_title(path: str | Path) -> str | None:
    """The document's embedded title, or ``None`` if it has none.

    Reads PDF/DOCX document metadata; other formats have no title. Callers
    typically fall back to the file name. Never raises -- a missing or unreadable
    title is just ``None``.
    """
    p = Path(path)
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            info = pypdf.PdfReader(str(p)).metadata
            title = info.title if info else None
        elif ext == ".docx":
            title = docx.Document(str(p)).core_properties.title
        else:
            return None
    except Exception:  # noqa: BLE001 -- best-effort metadata read; absence is fine
        return None
    title = (title or "").strip()
    return title or None
